#!/usr/bin/env python3
"""
extract.py — Structured content extraction from scientific papers.

Extracts metadata, sections, figures, tables, references, and associated
materials from PDF, DOCX, LaTeX, HTML, and Markdown documents into our
template engine format (content.yml + references.bib + figures/ + tables/).

Usage:
  python extract.py paper.pdf
  python extract.py paper.pdf --template scientific-paper/_template.yml -o output/
  python extract.py paper.docx --format docx
  python extract.py paper.tex --format latex
  python extract.py https://nature.com/articles/... --format html
  python extract.py paper.pdf --resolve-ids   # resolve DOI/PMID via APIs
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

import yaml

# ── Optional dependency imports (graceful fallback) ────────────────────────


def _try_import(module_name: str):
    """Attempt import, return None on failure."""
    try:
        return __import__(module_name)
    except ImportError:
        return None


fitz = _try_import("fitz")  # PyMuPDF
docx_mod = _try_import("docx")  # python-docx
bibtexparser = _try_import("bibtexparser")
bs4 = _try_import("bs4")
texsoup_mod = _try_import("TexSoup")
yaml_mod = yaml  # always available if pyyaml installed

# ── Sibling module imports ─────────────────────────────────────────────────

from .html import extract_html

# ── Format Detection ──────────────────────────────────────────────────────


def detect_format(path: str) -> str:
    """Auto-detect input format from file extension or URL."""
    if path.startswith("http://") or path.startswith("https://"):
        return "html"
    ext = Path(path).suffix.lower()
    return {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".tex": "latex",
        ".bib": "bibtex",
        ".md": "markdown",
        ".qmd": "markdown",
        ".html": "html",
        ".htm": "html",
    }.get(ext, "unknown")


# ── PDF Extraction ────────────────────────────────────────────────────────


def extract_pdf(path: str, output_dir: Path) -> dict[str, Any]:
    """Extract content from a PDF using PyMuPDF."""
    if fitz is None:
        print(
            "ERROR: PyMuPDF (fitz) not installed. Install with: pip install pymupdf",
            file=sys.stderr,
        )
        sys.exit(1)

    doc = fitz.open(path)
    result: dict[str, Any] = {
        "identifiers": {},
        "citation": {},
        "content": {},
        "figures": {},
        "references": {},
        "associated": {},
    }

    # ── Extract full text by page ──
    full_text = ""
    for page in doc:
        full_text += page.get_text("text") + "\n"

    # ── Extract DOI ──
    doi_match = re.search(r"(10\.\d{4,}/[^\s,;}\]]+)", full_text)
    if doi_match:
        result["identifiers"]["doi"] = doi_match.group(1).rstrip(".")

    # ── Extract title (heuristic: largest font on first page) ──
    first_page = doc[0]
    blocks = first_page.get_text("dict")["blocks"]
    title_candidates = []
    for block in blocks:
        if "lines" in block:
            for line in block["lines"]:
                for span in line["spans"]:
                    title_candidates.append(
                        {
                            "text": span["text"],
                            "size": span["size"],
                            "font": span["font"],
                            "y": span["origin"][1],
                        }
                    )

    # Sort by font size descending, take the largest text block near the top
    title_candidates.sort(key=lambda x: (-x["size"], x["y"]))
    if title_candidates:
        # Group consecutive spans with the same large font size as title
        max_size = title_candidates[0]["size"]
        title_parts = [
            s["text"] for s in title_candidates if s["size"] >= max_size * 0.95 and s["y"] < 300
        ]
        result["citation"]["title"] = " ".join(title_parts).strip()

    # ── Detect sections by font analysis ──
    sections = []
    current_section = None
    current_text = []

    # Heuristic: headers are bold or larger font, ALL CAPS, or follow patterns
    all_lines = full_text.split("\n")

    section_patterns = [
        r"^(Abstract|Introduction|Background|Methods|Materials and Methods|"
        r"Results|Discussion|Conclusion|Conclusions|Acknowledgments|Acknowledgements|"
        r"References|Bibliography|Supplementary|Supporting Information|"
        r"Data Availability|Code Availability|Author Contributions|"
        r"Competing Interests|Conflicts of Interest|Funding|Ethics)",
    ]
    combined_pattern = re.compile("|".join(section_patterns), re.IGNORECASE)

    for line in all_lines:
        stripped = line.strip()
        if not stripped:
            if current_section:
                current_text.append("")
            continue

        # Check if this line is a section header
        is_header = False
        if combined_pattern.match(stripped):
            is_header = True
        elif stripped.isupper() and 3 < len(stripped) < 80:
            is_header = True
        elif re.match(r"^\d+\.?\s+[A-Z]", stripped) and len(stripped) < 80:
            is_header = True

        if is_header:
            # Save previous section
            if current_section:
                sections.append(
                    {
                        "name": current_section,
                        "text": "\n".join(current_text).strip(),
                    }
                )
            current_section = stripped.rstrip(".")
            # Clean numbered prefix
            current_section = re.sub(r"^\d+\.?\s*", "", current_section)
            current_text = []
        elif current_section:
            current_text.append(stripped)

    # Save last section
    if current_section:
        sections.append(
            {
                "name": current_section,
                "text": "\n".join(current_text).strip(),
            }
        )

    # ── Extract abstract ──
    abstract_section = next((s for s in sections if s["name"].lower().startswith("abstract")), None)
    if abstract_section:
        result["content"]["abstract"] = abstract_section["text"]
        sections = [s for s in sections if s != abstract_section]
    else:
        # Try to find abstract in first page text
        abstract_match = re.search(
            r"(?:Abstract|ABSTRACT)[:\s]*(.+?)(?=\n\s*\n|\n[A-Z]{3,}|\n\d+\.\s)",
            full_text,
            re.DOTALL,
        )
        if abstract_match:
            result["content"]["abstract"] = abstract_match.group(1).strip()

    # ── Store sections ──
    result["content"]["sections"] = json.dumps(sections, indent=2)

    # ── Detect special sections ──
    for s in sections:
        name_lower = s["name"].lower()
        if "author contribution" in name_lower:
            result["content"]["author_contributions"] = s["text"]
        elif "acknowledgment" in name_lower or "acknowledgement" in name_lower:
            result["content"]["acknowledgments"] = s["text"]
        elif "funding" in name_lower:
            result["content"]["funding"] = s["text"]
        elif "competing interest" in name_lower or "conflict" in name_lower:
            result["content"]["competing_interests"] = s["text"]

    # ── Extract images/figures ──
    figures_dir = output_dir / "figures"
    figures_dir.mkdir(parents=True, exist_ok=True)
    figure_list = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        images = page.get_images(full=True)
        for img_idx, img_info in enumerate(images):
            xref = img_info[0]
            try:
                pix = fitz.Pixmap(doc, xref)
                if pix.n - pix.alpha > 3:  # CMYK
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                img_name = f"fig_p{page_num + 1}_{img_idx + 1}.png"
                img_path = figures_dir / img_name
                pix.save(str(img_path))
                figure_list.append(
                    {
                        "id": f"fig_p{page_num + 1}_{img_idx + 1}",
                        "caption": "",
                        "path": f"figures/{img_name}",
                        "page": page_num + 1,
                    }
                )
            except Exception:
                continue

    if figure_list:
        result["figures"]["figure_list"] = json.dumps(figure_list, indent=2)

    # ── Extract references section ──
    ref_section = next(
        (s for s in sections if s["name"].lower() in ("references", "bibliography")),
        None,
    )
    if ref_section:
        # Count references (heuristic: numbered entries)
        ref_lines = ref_section["text"].split("\n")
        ref_entries = [l for l in ref_lines if re.match(r"^\s*\[?\d+\]?\.?\s", l)]
        result["references"]["reference_count"] = (
            len(ref_entries) if ref_entries else len([l for l in ref_lines if l.strip()])
        )

    # ── Extract code/data availability ──
    for s in sections:
        text_lower = s["text"].lower()
        name_lower = s["name"].lower()
        if "code availability" in name_lower or ("code" in name_lower and "github" in text_lower):
            # Extract GitHub/GitLab URLs
            urls = re.findall(r"https?://(?:github|gitlab)\.com/[^\s,)]+", s["text"])
            if urls:
                result["associated"]["code_repositories"] = json.dumps(
                    [{"url": u, "description": ""} for u in urls], indent=2
                )
        if "data availability" in name_lower:
            # Extract accession numbers / database URLs
            accessions = re.findall(r"(?:GSE|GSM|SRP|SRX|E-MTAB-|PRJNA|syn)\d+", s["text"])
            urls = re.findall(r"https?://[^\s,)]+", s["text"])
            repos = []
            for acc in accessions:
                db = (
                    "GEO"
                    if acc.startswith("GS")
                    else "SRA"
                    if acc.startswith("SR")
                    else "ArrayExpress"
                    if acc.startswith("E-")
                    else "other"
                )
                repos.append({"accession": acc, "database": db, "url": "", "description": ""})
            for u in urls:
                if not any(r.get("url") == u for r in repos):
                    repos.append({"url": u, "accession": "", "database": "", "description": ""})
            if repos:
                result["associated"]["data_repositories"] = json.dumps(repos, indent=2)

    # ── Extract year ──
    year_match = re.search(r"(?:19|20)\d{2}", full_text[:2000])
    if year_match:
        result["citation"]["year"] = year_match.group()

    doc.close()
    return result


# ── DOCX Extraction ───────────────────────────────────────────────────────


def extract_docx(path: str, output_dir: Path) -> dict[str, Any]:
    """Extract content from a DOCX file using python-docx."""
    if docx_mod is None:
        print(
            "ERROR: python-docx not installed. Install with: pip install python-docx",
            file=sys.stderr,
        )
        sys.exit(1)

    from docx import Document

    doc = Document(path)

    result: dict[str, Any] = {
        "identifiers": {},
        "citation": {},
        "content": {},
        "figures": {},
        "references": {},
        "associated": {},
    }

    # ── Extract text by paragraphs ──
    sections = []
    current_section = None
    current_text = []

    for para in doc.paragraphs:
        style_name = para.style.name.lower() if para.style else ""
        text = para.text.strip()

        if not text:
            if current_section:
                current_text.append("")
            continue

        # Detect headers by style
        is_header = "heading" in style_name or "title" in style_name

        if is_header:
            if current_section:
                sections.append(
                    {
                        "name": current_section,
                        "text": "\n".join(current_text).strip(),
                    }
                )
            current_section = text
            current_text = []
        elif current_section is None:
            # Text before first heading — might be title/abstract
            if not result["citation"].get("title") and len(text) > 10:
                result["citation"]["title"] = text
        else:
            current_text.append(text)

    if current_section:
        sections.append({"name": current_section, "text": "\n".join(current_text).strip()})

    # ── Extract abstract ──
    abstract_section = next((s for s in sections if s["name"].lower().startswith("abstract")), None)
    if abstract_section:
        result["content"]["abstract"] = abstract_section["text"]
        sections = [s for s in sections if s != abstract_section]

    result["content"]["sections"] = json.dumps(sections, indent=2)

    # ── Extract tables ──
    tables_dir = output_dir / "tables"
    tables_dir.mkdir(parents=True, exist_ok=True)
    table_list = []

    for tbl_idx, table in enumerate(doc.tables, 1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            import csv

            tbl_name = f"table_{tbl_idx}.csv"
            tbl_path = tables_dir / tbl_name
            with open(tbl_path, "w", newline="") as f:
                writer = csv.writer(f)
                writer.writerows(rows)
            table_list.append(
                {
                    "id": f"tbl_{tbl_idx}",
                    "caption": "",
                    "path": f"tables/{tbl_name}",
                }
            )

    if table_list:
        result["figures"]["table_list"] = json.dumps(table_list, indent=2)

    return result


# ── LaTeX Extraction ──────────────────────────────────────────────────────


def extract_latex(path: str, output_dir: Path) -> dict[str, Any]:
    """Extract content from LaTeX source files."""
    result: dict[str, Any] = {
        "identifiers": {},
        "citation": {},
        "content": {},
        "figures": {},
        "references": {},
        "associated": {},
    }

    tex_content = Path(path).read_text(encoding="utf-8", errors="replace")

    # ── Title ──
    title_match = re.search(r"\\title\{([^}]+)\}", tex_content)
    if title_match:
        result["citation"]["title"] = title_match.group(1).strip()

    # ── Authors ──
    author_match = re.search(r"\\author\{(.+?)\}", tex_content, re.DOTALL)
    if author_match:
        result["citation"]["authors"] = re.sub(
            r"\\[a-z]+\{[^}]*\}", "", author_match.group(1)
        ).strip()

    # ── Abstract ──
    abstract_match = re.search(r"\\begin\{abstract\}(.+?)\\end\{abstract\}", tex_content, re.DOTALL)
    if abstract_match:
        result["content"]["abstract"] = abstract_match.group(1).strip()

    # ── Sections ──
    sections = []
    section_pattern = re.compile(r"\\(?:section|subsection|subsubsection)\*?\{([^}]+)\}", re.DOTALL)
    matches = list(section_pattern.finditer(tex_content))

    for i, match in enumerate(matches):
        section_name = match.group(1).strip()
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(tex_content)
        section_text = tex_content[start:end].strip()
        # Remove LaTeX commands roughly
        section_text = re.sub(r"\\[a-z]+\*?\{([^}]*)\}", r"\1", section_text)
        section_text = re.sub(r"\\[a-z]+", "", section_text)
        section_text = re.sub(r"[{}]", "", section_text)
        sections.append({"name": section_name, "text": section_text[:5000]})

    result["content"]["sections"] = json.dumps(sections, indent=2)

    # ── DOI ──
    doi_match = re.search(r"(?:doi|DOI)[:\s]*\{?(10\.\d{4,}/[^\s},]+)", tex_content)
    if doi_match:
        result["identifiers"]["doi"] = doi_match.group(1).rstrip(".")

    # ── Parse associated .bib file ──
    bib_match = re.search(r"\\bibliography\{([^}]+)\}", tex_content)
    if bib_match:
        bib_name = bib_match.group(1)
        if not bib_name.endswith(".bib"):
            bib_name += ".bib"
        bib_path = Path(path).parent / bib_name
        if bib_path.exists():
            result["references"]["references_bib"] = str(bib_path)
            if bibtexparser:
                with open(bib_path) as f:
                    bib_db = bibtexparser.load(f)
                result["references"]["reference_count"] = len(bib_db.entries)

    # ── Year ──
    year_match = re.search(r"\\date\{.*?((?:19|20)\d{2})", tex_content)
    if year_match:
        result["citation"]["year"] = year_match.group(1)

    return result


# ── Markdown Extraction ───────────────────────────────────────────────────


def extract_markdown(path: str, output_dir: Path) -> dict[str, Any]:
    """Extract content from Markdown/Quarto files."""
    result: dict[str, Any] = {
        "identifiers": {},
        "citation": {},
        "content": {},
        "figures": {},
        "references": {},
        "associated": {},
    }

    content = Path(path).read_text(encoding="utf-8", errors="replace")

    # ── Parse YAML frontmatter ──
    if content.startswith("---"):
        parts = content.split("---", 2)
        if len(parts) >= 3:
            try:
                frontmatter = yaml.safe_load(parts[1])
                if isinstance(frontmatter, dict):
                    if "title" in frontmatter:
                        result["citation"]["title"] = frontmatter["title"]
                    if "author" in frontmatter:
                        authors = frontmatter["author"]
                        if isinstance(authors, list):
                            result["citation"]["authors"] = ", ".join(
                                a.get("name", str(a)) if isinstance(a, dict) else str(a)
                                for a in authors
                            )
                        else:
                            result["citation"]["authors"] = str(authors)
                    if "date" in frontmatter:
                        year_match = re.search(r"((?:19|20)\d{2})", str(frontmatter["date"]))
                        if year_match:
                            result["citation"]["year"] = year_match.group(1)
                    if "abstract" in frontmatter:
                        result["content"]["abstract"] = str(frontmatter["abstract"])
                    if "doi" in frontmatter:
                        result["identifiers"]["doi"] = frontmatter["doi"]
                    if "bibliography" in frontmatter:
                        result["references"]["references_bib"] = frontmatter["bibliography"]
                content = parts[2]
            except yaml.YAMLError:
                pass

    # ── Extract sections by headers ──
    sections = []
    current_section = None
    current_text = []

    for line in content.split("\n"):
        header_match = re.match(r"^(#{1,3})\s+(.+)", line)
        if header_match:
            if current_section:
                sections.append({"name": current_section, "text": "\n".join(current_text).strip()})
            current_section = header_match.group(2).strip()
            current_text = []
        elif current_section:
            current_text.append(line)

    if current_section:
        sections.append({"name": current_section, "text": "\n".join(current_text).strip()})

    result["content"]["sections"] = json.dumps(sections, indent=2)

    return result


# ── Content → YAML Export ─────────────────────────────────────────────────


def export_content_yml(data: dict[str, Any], output_path: Path):
    """Export extracted data to content.yml format."""
    # Clean up: remove empty values
    cleaned = {}
    for section_id, fields in data.items():
        if isinstance(fields, dict):
            non_empty = {k: v for k, v in fields.items() if v}
            if non_empty:
                cleaned[section_id] = non_empty

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with open(output_path, "w") as f:
        f.write("# Extracted content — auto-generated\n")
        f.write(f"# Source: {data.get('_source', 'unknown')}\n\n")
        # Remove internal metadata before export
        cleaned.pop("_source", None)
        yaml.dump(
            cleaned,
            f,
            default_flow_style=False,
            allow_unicode=True,
            width=120,
            sort_keys=False,
        )

    print(f"✓ Content exported to {output_path}", file=sys.stderr)


# ── ID Resolution ─────────────────────────────────────────────────────────


def resolve_ids(data: dict[str, Any]):
    """Resolve DOI/PMID via CrossRef and PubMed APIs."""
    doi = data.get("identifiers", {}).get("doi")
    if not doi:
        return

    # ── CrossRef: DOI → metadata ──
    try:
        import urllib.request

        url = f"https://api.crossref.org/works/{doi}"
        req = urllib.request.Request(
            url,
            headers={"User-Agent": "Cytognosis/1.0 (mailto:mohammadi@cytognosis.org)"},
        )
        with urllib.request.urlopen(req, timeout=10) as resp:
            cr_data = json.loads(resp.read())
            msg = cr_data.get("message", {})

            # Fill missing fields
            if not data["citation"].get("title"):
                titles = msg.get("title", [])
                if titles:
                    data["citation"]["title"] = titles[0]
            if not data["citation"].get("journal"):
                containers = msg.get("container-title", [])
                if containers:
                    data["citation"]["journal"] = containers[0]
            if not data["citation"].get("year"):
                issued = msg.get("issued", {}).get("date-parts", [[]])
                if issued and issued[0]:
                    data["citation"]["year"] = str(issued[0][0])
            if not data["citation"].get("volume"):
                data["citation"]["volume"] = msg.get("volume", "")
            if not data["citation"].get("pages"):
                data["citation"]["pages"] = msg.get("page", "")
            if not data["citation"].get("publisher"):
                data["citation"]["publisher"] = msg.get("publisher", "")

            # Authors
            if not data["citation"].get("authors"):
                authors = msg.get("author", [])
                author_strs = [
                    f"{a.get('given', '')} {a.get('family', '')}".strip() for a in authors
                ]
                data["citation"]["authors"] = ", ".join(author_strs)

            print(f"✓ CrossRef metadata resolved for DOI: {doi}", file=sys.stderr)
    except Exception as e:
        print(f"⚠ CrossRef resolution failed: {e}", file=sys.stderr)

    # ── PubMed: DOI → PMID ──
    if not data.get("identifiers", {}).get("pmid"):
        try:
            url = f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi?db=pubmed&term={doi}[doi]&retmode=json"
            with urllib.request.urlopen(url, timeout=10) as resp:
                pm_data = json.loads(resp.read())
                id_list = pm_data.get("esearchresult", {}).get("idlist", [])
                if id_list:
                    data["identifiers"]["pmid"] = id_list[0]
                    print(f"✓ PubMed ID resolved: {id_list[0]}", file=sys.stderr)
        except Exception as e:
            print(f"⚠ PubMed resolution failed: {e}", file=sys.stderr)


# ── Main CLI ──────────────────────────────────────────────────────────────


def main():
    parser = argparse.ArgumentParser(
        description="Extract structured content from scientific papers",
    )
    parser.add_argument("input", help="Input file path or URL")
    parser.add_argument("-f", "--format", help="Force input format (pdf/docx/latex/html/markdown)")
    parser.add_argument("-o", "--output-dir", default=".", help="Output directory (default: .)")
    parser.add_argument("--template", help="Path to _template.yml for guided extraction")
    parser.add_argument(
        "--resolve-ids",
        action="store_true",
        help="Resolve DOI/PMID via CrossRef/PubMed APIs",
    )
    parser.add_argument(
        "--backend",
        choices=["auto", "pymupdf", "grobid", "docling"],
        default="auto",
        help="PDF extraction backend (default: auto)",
    )
    parser.add_argument(
        "--enrich",
        action="store_true",
        help="Enrich with Semantic Scholar API",
    )
    parser.add_argument(
        "--ner",
        action="store_true",
        help="Extract biomedical entities via scispaCy",
    )
    parser.add_argument(
        "--ocr",
        action="store_true",
        help="Force OCR pre-processing for scanned PDFs",
    )
    args = parser.parse_args()

    fmt = args.format or detect_format(args.input)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"→ Extracting from: {args.input} (format: {fmt})", file=sys.stderr)

    extractors = {
        "pdf": extract_pdf,
        "docx": extract_docx,
        "latex": extract_latex,
        "markdown": extract_markdown,
        "html": extract_html,
    }

    if fmt not in extractors:
        print(
            f"ERROR: Unsupported format '{fmt}'. Supported: {', '.join(extractors.keys())}",
            file=sys.stderr,
        )
        sys.exit(1)

    # ── OCR pre-processing for scanned PDFs ──
    if fmt == "pdf" and args.ocr:
        from .ocr import is_scanned_pdf
        from .ocr import ocr_pipeline as run_ocr

        if is_scanned_pdf(args.input):
            ocr_text = run_ocr(args.input, output_dir)
            if ocr_text:
                # Save OCR text and extract from that
                ocr_md = output_dir / "ocr_output.md"
                ocr_md.write_text(ocr_text, encoding="utf-8")
                print("→ Using OCR output for extraction", file=sys.stderr)
                data = extract_markdown(str(ocr_md), output_dir)
                data["_source"] = args.input
                data["_ocr"] = True

                if args.resolve_ids:
                    resolve_ids(data)
                if args.enrich or args.ner:
                    from .enrichment import enrich_content

                    data = enrich_content(
                        data,
                        output_dir,
                        enable_s2=args.enrich,
                        enable_ner=args.ner,
                    )
                export_content_yml(data, output_dir / "content.yml")
                return

    # ── Backend selection for PDF ──
    if fmt == "pdf" and args.backend != "pymupdf":
        from .grobid import extract_pdf_grobid, is_grobid_available

        if args.backend == "grobid" or (args.backend == "auto" and is_grobid_available()):
            if is_grobid_available():
                print("→ Using GROBID backend", file=sys.stderr)
                data = extract_pdf_grobid(args.input, output_dir)
            elif args.backend == "grobid":
                print(
                    "ERROR: GROBID not available. Start with: "
                    "docker run --rm -p 8070:8070 lfoppiano/grobid:0.8.2",
                    file=sys.stderr,
                )
                sys.exit(1)
            else:
                print(
                    "⚠ GROBID not available, falling back to PyMuPDF",
                    file=sys.stderr,
                )
                data = extract_pdf(args.input, output_dir)
        else:
            data = extract_pdf(args.input, output_dir)
    else:
        data = extractors[fmt](args.input, output_dir)

    data["_source"] = args.input

    if args.resolve_ids:
        resolve_ids(data)

    # ── Post-extraction enrichment ──
    if args.enrich or args.ner:
        from .enrichment import enrich_content

        data = enrich_content(
            data,
            output_dir,
            enable_s2=args.enrich,
            enable_ner=args.ner,
        )

    export_content_yml(data, output_dir / "content.yml")


if __name__ == "__main__":
    main()
